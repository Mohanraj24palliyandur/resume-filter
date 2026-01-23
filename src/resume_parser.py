import fitz  # PyMuPDF
import docx
import os
from typing import List, Dict
from pyresparser import ResumeParser
import pdfplumber
from docx import Document
import re


SKILLS = [
    "python", "java", "nlp", "machine learning", "deep learning",
    "pandas", "numpy", "sql", "flask", "django", "streamlit",
    "javascript", "react", "html", "css", "tensorflow", "pytorch"
]


class ResumeParser:
    """
    A class to parse resumes from PDF and DOCX formats and extract text content.
    """

    def __init__(self, resume_dir: str = "data/resumes/"):
        """
        Initialize the ResumeParser with the directory containing resumes.

        Args:
            resume_dir (str): Path to the directory containing resume files.
        """
        self.resume_dir = resume_dir

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.

        Args:
            file_path (str): Path to the PDF file.

        Returns:
            str: Extracted text from the PDF.
        """
        text = ""
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
        except Exception as e:
            print(f"Error extracting text from PDF {file_path}: {e}")
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_path (str): Path to the DOCX file.

        Returns:
            str: Extracted text from the DOCX file.
        """
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {e}")
        return text

    def parse_resume(self, file_path: str) -> str:
        """
        Parse a resume file (PDF or DOCX) and extract text.

        Args:
            file_path (str): Path to the resume file.

        Returns:
            str: Extracted text from the resume.
        """
        if file_path.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path}")

    def get_all_resumes(self) -> Dict[str, str]:
        """
        Get all resumes from the resume directory and extract their text.

        Returns:
            Dict[str, str]: Dictionary with filename as key and extracted text as value.
        """
        resumes = {}
        if not os.path.exists(self.resume_dir):
            print(f"Resume directory {self.resume_dir} does not exist.")
            return resumes

        for filename in os.listdir(self.resume_dir):
            if filename.lower().endswith(('.pdf', '.docx')):
                file_path = os.path.join(self.resume_dir, filename)
                try:
                    text = self.parse_resume(file_path)
                    resumes[filename] = text
                except Exception as e:
                    print(f"Error parsing resume {filename}: {e}")

        return resumes

    def parse(self, file_path):
        # Extract full_text
        if file_path.lower().endswith('.pdf'):
            with pdfplumber.open(file_path) as pdf:
                full_text = " ".join(page.extract_text() or '' for page in pdf.pages).strip()
        elif file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            full_text = " ".join(para.text for para in doc.paragraphs).strip()
        else:
            full_text = ''
        
        # Extract using heuristics
        name = self.extract_name(full_text)
        experience_years = self.extract_experience_years(full_text)
        skills = self.extract_skills(full_text)
        
        # Try pyresparser for overrides
        try:
            data = ResumeParser(file_path).get_extracted_data()
            if data.get('name') and data['name'] != 'Not extracted':
                name = data['name']
            if data.get('total_experience'):
                experience_years = data['total_experience']
            if data.get('skills'):
                skills = data['skills']
        except:
            pass
        
        return {
            'full_text': full_text,
            'name': name,
            'email': '',
            'skills': skills,
            'experience_years': experience_years
        }

    def extract_name(self, text: str) -> str:
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if not lines:
            return "Not extracted"
        first_line = lines[0]
        if "@" in first_line or any(char.isdigit() for char in first_line):
            return "Not extracted"
        if len(first_line.split()) <= 4:
            return first_line.title()
        return "Not extracted"

    def extract_experience_years(self, text: str) -> int:
        patterns = [
            r'(\d+)\+?\s+years?',
            r'(\d+)\s+yrs',
            r'experience:?\s*(\d+)\s+years?'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return int(match.group(1))
        return 0

    def extract_skills(self, text: str) -> list:
        text_lower = text.lower()
        found = [skill for skill in SKILLS if skill in text_lower]
        return found[:5]


if __name__ == "__main__":
    # Example usage
    parser = ResumeParser()
    resumes = parser.get_all_resumes()
    print(f"Parsed {len(resumes)} resumes")
    for filename, text in resumes.items():
        print(f"{filename}: {len(text)} characters")